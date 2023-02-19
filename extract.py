""" This module uses 'python-docx' package to analyze a Word Document ('.docx')
    This module can be imported in other programs but the preffered way is to run it 
    as a Script from the command line like this:
    
    $ python extract.py <name-of-word-document-file> <step>
    
    where <step> is approximately how many words the program will analyze in each loop. 
    
    For example, let's say there are 50 words in a word document.
      If we run this program as a script on this word document and
      provide <step> of 10 then the program will analyze 
      10 words for each iteration
      i.e. 0 - 10 , 10 - 20, 20 - 30, 30 - 40, 40 - 50

    The basic function of the program is to analyze the word document and 
    collect bold, italic and underlined words from it and then 
    after analyzing write these collected words at the end of the word document.

    So the program first reads the word document, collects bold, italic 
    and underlined words from it and then writes the collected words at
    the very end of the same word document 

    Copyright 2023 Kashaan Mahmood

    License: MIT License
             https://opensource.org/license/mit/
    """


from math import ceil
from docx import Document
from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# global variables
total_words = 0
start = 0
end = 0
step = 20
wordsList = ""
FROM = ""


# calculate total words in docx
def get_total_words(docxFile):
    document = Document(docxFile)
    total = 0
    for p in document.paragraphs:
        for run in p.runs:
            total += len(run.text.split())
    return total


unwanted_characters = [
    '"',
    "'",
    "’",
    "“",
    ":",
    "\n",
    "-",
    "— — ",
    "—",
    ".",
    ",",
    ";",
    "!",
    "?",
]


def remove_unwanted(words):
    """remove unwanted characters from analyzed output"""
    for i in unwanted_characters:
        if i in words:
            words = words.replace(i, "")
    return words


def analyze(docxFile):
    """analyze the docx file and collect bold, italicized, and underlined words from it
    and return a `collect` object these selected words
    """

    global start, end, step, total_words

    total_words = get_total_words(docxFile)

    if total_words - end > 0:
        end = start + step
    else:
        end = total_words

    document = Document(docxFile)

    collect = [
        {"b": []},
        {"i": []},
        {"u": []},
        {"bi": []},
        {"bu": []},
        {"iu": []},
        {"biu": []},
    ]

    # checks
    checkRange = 0

    for p in document.paragraphs:
        for run in p.runs:
            checkRange += len(run.text.split())

            if checkRange > start and checkRange < end:
                if run.bold and run.italic and run.underline:
                    filtered_text = remove_unwanted(run.text)
                    collect[6]["biu"].append(filtered_text)

                elif run.bold and run.italic:
                    filtered_text = remove_unwanted(run.text)
                    collect[3]["bi"].append(filtered_text)

                elif run.bold and run.underline:
                    filtered_text = remove_unwanted(run.text)
                    collect[4]["bu"].append(filtered_text)

                elif run.italic and run.underline:
                    filtered_text = remove_unwanted(run.text)
                    collect[5]["iu"].append(filtered_text)

                elif run.bold:
                    filtered_text = remove_unwanted(run.text)
                    collect[0]["b"].append(filtered_text)

                elif run.italic:
                    filtered_text = remove_unwanted(run.text)
                    collect[1]["i"].append(filtered_text)

                elif run.underline:
                    filtered_text = remove_unwanted(run.text)
                    collect[2]["u"].append(filtered_text)

            elif checkRange >= end:
                return collect


def write_data(docxFile, data):
    """gets the `collect` variable as 'data' argument from analyze()
    and reads and appends the 'data' to end of docx file
    """

    global start, end, wordsList, FROM

    document = Document(docxFile)

    def save_document():
        document.save(docxFile)
        return "saved"

    # checks
    checkEnd = 0
    diff = 0

    for p in document.paragraphs:
        for run in p.runs:
            checkEnd += len(run.text.split())
            diff = end - checkEnd

            if checkEnd == end or (diff < 1 and end > checkEnd):

                def add_words(key):
                    global wordsList

                    for word in words[key]:
                        if not word == "" and not word == " ":
                            if key == "b":
                                category = "Bold Words:-"
                                if len(wordsList) == 0 or category not in wordsList:
                                    wordsList = wordsList + f"{category}\n{word}"
                                else:
                                    wordsList = wordsList + f", {word}"

                            elif key == "i":
                                category = "Italicized Words:-"
                                if len(wordsList) == 0 or category not in wordsList:
                                    wordsList = wordsList + f"\n\n{category}\n{word}"
                                else:
                                    wordsList = wordsList + f", {word}"

                            elif key == "u":
                                category = "Underlined  Words:-"
                                if len(wordsList) == 0 or category not in wordsList:
                                    wordsList = wordsList + f"\n\n{category}\n{word}"
                                else:
                                    wordsList = wordsList + f", {word}"

                            elif key == "bi":
                                category = "Bold & Italicized Words:-"
                                if len(wordsList) == 0 or category not in wordsList:
                                    wordsList = wordsList + f"\n\n{category}\n{word}"
                                else:
                                    wordsList = wordsList + f", {word}"

                            elif key == "bu":
                                category = "Bold & Underlined Words:-"
                                if len(wordsList) == 0 or category not in wordsList:
                                    wordsList = wordsList + f"\n\n{category}\n{word}"
                                else:
                                    wordsList = wordsList + f", {word}"

                            elif key == "iu":
                                category = "Italicized & Underlined Words:-"
                                if len(wordsList) == 0 or category not in wordsList:
                                    wordsList = wordsList + f"\n\n{category}\n{word}"
                                else:
                                    wordsList = wordsList + f", {word}"

                            elif key == "biu":
                                category = "Bold & Italicized & Underlined Words:-"
                                if len(wordsList) == 0 or category not in wordsList:
                                    wordsList = wordsList + f"\n\n{category}\n{word}"
                                else:
                                    wordsList = wordsList + f", {word}"

                title_p = document.add_paragraph(
                    f"\n========== Extracted Words - Range [{start} - {end}] ==========\n"
                )
                title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                write_p = document.add_paragraph()

                if not FROM:  # Start of file
                    write_p.add_run("From:\n")
                    write_p.add_run("START OF FILE \n\n")

                else:
                    write_p.add_run("From:\n")
                    write_p.add_run(f"{FROM}\n\n")

                write_p.add_run("To:\n")
                write_p.add_run(f"'...{run.text}'\n\n")

                FROM = f"'{run.text}...'"

                for words in data:
                    if words.__contains__("b") and words["b"]:
                        add_words("b")

                    elif words.__contains__("i") and words["i"]:
                        add_words("i")

                    elif words.__contains__("u") and words["u"]:
                        add_words("u")

                    elif words.__contains__("bi") and words["bi"]:
                        add_words("bi")

                    elif words.__contains__("bu") and words["bu"]:
                        add_words("bu")

                    elif words.__contains__("iu") and words["iu"]:
                        add_words("iu")

                    elif words.__contains__("biu") and words["biu"]:
                        add_words("biu")

                write_p.add_run(f"{wordsList}")

                ending_p = document.add_paragraph("\n===================\n")
                ending_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                save_document()
                break

            elif checkEnd > end:
                return

        # reset wordsList
        wordsList = ""


# function calls inside main()


def main():
    global start, end, step, total_words

    if total_words == 0:
        total_words = get_total_words(docx)

    # number of times the program will read and write to docx
    cycles = ceil(total_words / step)

    for i in range(cycles):
        data = analyze(docx)
        write_data(docx, data)
        start = end
        end += step


if __name__ == "__main__":
    from sys import argv
    import time

    # get docx file
    docx = argv[1]
    step = int(argv[2])
    print(f"Started at {time.strftime('%X')}...")
    # calling main()
    main()
    print(f"Finished at {time.strftime('%X')}...")
